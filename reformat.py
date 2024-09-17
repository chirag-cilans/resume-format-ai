import fitz
import io
import os
import pdfplumber
import re
import streamlit as st
import subprocess
import tempfile
import win32com.client
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx2pdf import convert
from openai import OpenAI
import pythoncom

# Constants
FONT_NAME = "Times New Roman"
LOGO_PATH = "kyralogo.png"
CONTACT_INFO = """3673 Coolidge Ct.,
Tallahassee, FL 32311
Phone: (850) 459-5854
Email: vpatel@KyraSolutions.com"""


# Initialize the OpenAI client
def initialize_openai_client(api_key):
    return OpenAI(api_key=api_key)


def extract_content_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text


def clean_text(text):
    text = re.sub(r"\s+", " ", text)
    text = "".join(char for char in text if char.isprintable() or char in ["\n", "\t"])
    return text.strip()


def wrap_keywords_in_b_tags(text, keywords):
    # Sort keywords by length in descending order to avoid partial replacements
    keywords = sorted(keywords, key=len, reverse=True)

    for keyword in keywords:
        # Use re.sub with a lambda to replace the matched keyword with <b>wrapped keyword</b>, ignoring case
        text = re.sub(
            f"(?i)({re.escape(keyword)})",
            lambda match: f"<b>{match.group(0)}</b>",
            text,
        )

    return text


def reformat_resume(content, keywords):

    formatted_resume = """
Chitramalini Sathiraju
Business Analyst

Please Note: Chitramalini is currently based in Folsom, CA, and is ready to relocate to Tallahassee upon
receiving an offer.
PROFESSIONAL SUMMARY
 Multi-Certified Business Analyst/Quality Assurance professional with 16 years of experience
in the IT Industry.
 Experience working with PEGA Systems, Client/Server, Web, and ERP, Salesforce - CRM, and
Guidewire applications.
 Experience working with the State of West Virginia and State of Maryland.
 Proficient in complex, Agile, Scrum, V, and Waterfall methodologies.
 Experience in Requirements Capture, Requirements Categorization, Requirements
Prioritization, Requirements Documentation, Business Rules, stakeholder analysis, Gap
analysis and Impact analysis.
 Good experience in Interface or API testing with Postman, SOAP UI tools, batch testing with Win
SCP, and Notepad++, creating, modifying, and enhancing manual Test cases, and Data migration
applications from legacy systems to As-IS systems.
 Experience in entire QA life cycle activities, testing concepts, Bug Tracking, Root Cause
Analysis, Defect Analysis, Project Documentation, Designing, Developing, and Generating
Reports, interacting with business analysts, developers, and technical support teams and helping
them in baselining the requirement specifications, maintaining support documents, QA Sign-off
Documents, maintaining Status Reports, creating Requirements Traceability Matrix to ensure
comprehensive test coverage of requirements, preparation of Test plans, Test procedures- and
Manual Test Scenarios.
 Well understanding of Software Development Life Cycle (SDLC), Software Testing Life Cycle
(STLC), Requirements Life Cycle, and Bug Life Cycle methodologies.
 Experience in developing SQL queries for backend database testing, and testing applications
migrating from .Net platform to JAVA platform.
TECHNICAL SKILLS
Operating Systems Windows 95, 98, XP/2000/NT, MS-DOS, UNIX
Languages C, C++, C#, JAVA, PEGA
GUI VB6.0
ERP Mainframes, SAP
Database Oracle10, DB Visualizer, Data Trek
Web Technologies VBScript, JavaScript, HTML, VB.Net, ASP. Net, XML, PEGA
Package MS-Office
Testing Tools QTP12, Quality Center12, LoadRunner11.52, Selenium-Web-Driver, JIRA,

SQUIDS, RALLY, ALM, TDD, Microsoft, AZURE, TFS

CRM Guidewire, Salesforce, Business Analytics
Analytics Tools Rapid Minor
Diagramming Tools Axure RP 10
EDUCATION QUALIFICATION
 MSc. IT/Vinayaka Mission Deemed University, 2005
 BSc Computer Science/Osmania University, India,1998


CERTIFICATION/TRAINING
 GAQM Certified Scrum Master (CSM)®
 SAFe 4 Scrum Master Certificate
 Safe 4 Agilist Certificate
 ISEB ISTQB Certified at the foundation level
 Expert Rating in Software Testing
 Brain bench in Software Testing
 MIT No Code AI and Machine Learning Building Data Science Solutions Certificate

WORK HISTORY
Date: 01/23 - 07/24 (19 Months)
Company: State of California, Franchise Tax Board, Folsom, CA
Title: Business Analyst/Quality Analyst
Tools and Technologies: PEGA, Axure RP 10, DATABASICS, Manual Testing, Report testing API
testing, Microsoft SQL Server, SoapUI, Postman V9.15.2 WINDOWS 10, Google Chrome, Cisco
VPN WINSCP5.15.3, Notepad++, Azure, TFS, Teams, TOSCO
Roles and Responsibilities
 Requirements Capture - Engage with business owners to capture Epics and User Stories.
Facilitate workshops.
 Requirements Categorization - Refine and categorize requirements into Functional &amp; Non-
functional requirements.
 Requirements Prioritization - Use MoSCoW methodology to prioritize requirements. Work
with business change owners and IT suppliers to ensure that requirements are clearly understood
and necessary information is available.
 Requirements Documentation - Document requirements catalog, functional specifications,
Use Case Analysis, and UML artifacts.
 Responsible for Business Rules - Revise and create new business rules for approvals,
delegations, and hard-charging and conducted stakeholder analysis and identified and resolved
Impact/relationship/dependencies/constraints.
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, and User Acceptance Testing, with developers for
requirement development.
 Tracked Change requests, defects, and test coverage, and published Defect Status Reports using
TFS.
 Ensured strategic objectives, strategy definition, and target models are clear, documented, and
collectively agreed upon.
 Tested PEGA environment and customized the PEGA application as per the client&#39;s requirements,
Integration of PEGA systems with Standalone Forms (SAF), Governance Process Rules Engine
(Rules), Case Management (CM), Return Analysis (RA), Return Verification (RV), Collections,
Notices, CX applications, iCapture, ETS Database and execute the Workflows.
 Performed agile activities using the DATABASICS tool.
 Managed client information for business communications using Microsoft Dynamics 365.

Date: 06/22 - 12/22 (07 Months)
Company: United Health Group, Minneapolis, MN
Client: Child Support State of West Virginia
Title: Analyst/Tester
Tools and Technologies: Sales Force.com, Force.com platform, Service Cloud, Data Loader, e-
commerce, ALM, Rally, TDD, Deltek, Manual Testing, Report testing API testing, Microsoft SQL
Server, SoapUI, Postman V9.15.2 WINDOWS 10, Google Chrome, Cisco VPN WINSCP5.15.3,
Notepad++, FACETS TriZetto
Roles and Responsibilities
 Validated Data after data migration, KPI reports data with AS-IS system, the response and
result XMLs, JSON script issues in Postman.
 Involved in testing, Batch testing, API interfaces, and Web Services, analyzing migration issues
in Agile methodologies.
 Trained Maryland county users in using our applications.
 Mentored the Tester, arranging team meetings, getting updates from each person, forwarding
weekly reports to all the team members and managers, interacting with stakeholders in the initial
stage until handed over automation tester, and forwarding release reports to appropriate people.
 Tracked requirements, and test coverage using VersionOne tool, defects, and published Defect
Status Reports using JIRA.
 Created SIT and UAT test cases for functional testing and regression testing.
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, and User Acceptance Testing.
 Tested Child Support (CSMS), etc modules with sanity flows with migrated data for releasing it to
West Virginia counties.
 Prepared inbound files and Outbound files.
 Performed agile activities using the Deltek tool.
 Managed client information for business communications using Microsoft Dynamics 365.
 Supported county users when they were facing issues.


Date: 01/19 - 05/22 (41 Months)
Company: MDThink, Linthicum Heights, MD
Client: CJAMS State of Maryland
Title: Analyst/Tester
Tools and Technologies: Sales Force.com, Force.com platform, Service Cloud, Data Loader,
ecommerce, JIRA, VersionOne, Digital.ai Agility, Manual Testing, Report testing API testing,
DBeaver /Postgres admin, SoapUI, Postman V9.15.2 WINDOWS 10, Google Chrome, Open Vpn,
Zscaler, WINSCP5.15.3, Notepad++
Roles and Responsibilities
 Tested Child Welfare (CWS), Child Support (CSMS), etc modules in Agile methodology with
sanity flows with migrated data for releasing it to Maryland counties.
 Validated Data after data migration, KPI reports data with AS-IS system, the response and
result XMLs, JSON script issues in Postman.
 Mentored the Tester, arranging team meetings, getting updates from each person, forwarding
weekly reports to all the team members and managers, interacting with stakeholders in the initial
stage until handed over automation tester, and forwarding release reports to appropriate people.
 Tracked requirements, and test coverage using the VersionOne tool, defects, and published Defect
Status Reports using JIRA.
 Created SIT and UAT test cases for functional testing and regression testing.
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, and User Acceptance Testing.
 Prepared inbound files and Outbound files.
 Involved in testing, Batch testing, API interfaces, and Web Services, analyzing migration issues.
 Performed agile activities using the Digital.ai Agility tool.
 Managed client information for business communications using Microsoft Dynamics 365.
 Trained Maryland county users in using our applications.
 Supported county users when they are facing issues.


Date: 03/17 - 12/18 (22 Months)
Company: CORVUS, Greater Pittsburgh Area
Client: Best Nomos
Title: Senior QA consultant/Team Lead
Tools and Technologies: eCommerce platform, Test Tracking Tool, PTWeb, DOORS, Manual
Testing, Mac OS, Safari, Firefox, ALM and RALLY
Roles and Responsibilities
 Responsible for the preparation of the Test plan, test approach, Business Requirement
document, and deliverables.
 Mentored Tester, arranging team meetings, getting updates from each person, forwarding weekly
reports to all the team members and managers, interacting with stakeholders in the initial stage
until handed over automation tester, and forwarding release reports to appropriate people.
 Tracked requirements using the DOORS tool, the defects, and published Defect Status Reports
using PTWeb, test coverage using RTM, TTT.
 Created test cases for functional testing and regression testing, user acceptance testing using
Excel.
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, and User Acceptance Testing.
 Tested CORVUS14, CORVUS Beam Utilities14, CORVUS Configuration Manager14, Corvus
User Manager14, and CORVUS Structure Set Editor modules for Forward planning, Inverse
planning, Treatment planning process in the integration with beam parameters and dose
distribution in modules in Waterfall methodology.

Date: 01/15 - 02/17 (26 Months)
Company: SOPHIA, Englewood, CO
Client: Tele Tech
Title: Senior QA consultant/Team Lead
Tools and Technologies: .Net, ALM, and RALLY
Roles and Responsibilities
 Prepared the Test plan, test approach, and deliverables.
 Created, modified, and enhanced manual Test cases.
 Created test cases for functional testing and regression testing, workflow diagrams for
development references, and data mapping.
 Tested Cards and Payments module for the Card Determination process in the integration with
Customer Management and Payment Type modules in Waterfall methodology.
 Guided Tester, arranging team meetings, getting updates from each person, forwarding weekly
reports to all the team members and managers, interacting with stakeholders in the initial stage
until handed over automation tester, and forwarding release reports to appropriate people.
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing and Regression
Testing.
 Tracked the defects, published Defect Status Reports using ALM, test coverage using RALLY.

Date: 05/13 - 12/14 (20 Months)
Company: Enterprise Pricing Request System (ePRS), TN
Client: FedEx
Title: Technical Consultant/Team Lead
Tools and Technologies: E-commerce platform, Java, JSP, ASP servlets, QC12.0
Roles and Responsibilities
 Analyzed the requirements from the Business Requirements, Functional Requirements, and
High-level design documents, and test results to make sure requirements have been met and
existing system requirement still holds.
 Guided Tester, arranging team meetings, getting updates from each person, forwarding weekly
reports to all the team members and managers, interacting with stakeholders in the initial stage
until handed over to the Client-side Test Manager, and forwarding release reports to appropriate
people.
 Used QC for test management i.e. Defect management, Bug Reporting, and Bug tracking.
 Coordinated and sent daily status to Project Managers.
 Created, modified, and enhanced manual Test cases in Agile methodology.
Date: 08/12 - 04/13 (09 Months)
Company: Davis Negative Files, San Jose, CA
Client: PayPal
Title: Team Lead
Tools and Technologies: Java, UNIX, JIRA, Java, UNIX, Data Trek
Roles and Responsibilities
 Developed test plans, test cases, and test scenarios to map Integration Testing, System Testing,
and User Acceptance Testing to test business requirements and design documents according to
the BRD&#39;s &amp; FRDs.
 Monitored and tracked priority defects on a day-to-day basis, wrote positive and negative test
cases, improved traceability, reduced project risk, and increased the quality of the application in
Agile methodology.
 Responsible for managing the change request process related to testing.
 Involved in execution of SIT and UAT, Cycle 1, Cycle 2 &amp; Cycle 3 Testing, daily meetings to
review the progress of testing as well as discuss the status of defects with offshore people in the
USA.
 Responsible for Coordination and communication between US Clients and SA testing Effort.
 Used HP Quality Center (QC) to Report bugs using the defect tracking system and verify fixes for
reserving, executing test cases, defect tracking, and complete test management, to get the status
reports of all the test cases &amp;test scripts that are executed during the testing process, re-testing of
defects corrected and reported.

Date: 10/11 - 07/12 (10 Months)
Company: Retail Staffing Administration, Atlanta, GA
Client: Home Depot
Title: QA Consultant
Tools and Technologies: E-commerce platform, Java, .Net and SOA, QC10, DB Visualizer and QTP
Roles and Responsibilities
 Developed test cases, and test scenarios to map Integration Testing, System Testing, and User
Acceptance Testing to test business requirements and design documents according to the BRD&#39;s
FRDs.
 Used HP Quality Center (QC) to Report bugs using the defect tracking system and verify fixes for
reserving, executing test cases, defect tracking, and complete test management, to get the status
reports of all the test cases &amp; test scripts that are executed during the testing process, re-testing of
defects corrected and reported.
 Responsible for managing the change request process related to testing.
 Involved in execution of SIT and UAT, Cycle 1, Cycle 2 &amp; Cycle 3 Testing, daily meetings to
review the progress of testing as well as discuss the status of defects with offshore people in the
USA.
 Responsible for Coordination and communication between US Clients and SA testing Effort.
 Monitored and tracked priority defects on a day-to-day basis, wrote positive and negative test
cases, improved traceability, reduced project risk, and increased the quality of the application in
Agile methodology.

"""

    unformatted_resume = """
    Chitramalini Sathiraju
Email: chitra.s@selectiva.io
Phone: 412-652-2565| Alt: 408-669-5160
PROFESSIONAL SUMMARY:
 Multi-Certified Business Analyst/Quality Assurance professional with 16 years of experience in
the IT Industry.
 Worked with PEGA Systems, Client/Server, Web, and ERP, Salesforce - CRM, and Guidewire
applications.
 Well understanding of Software Development Life Cycle (SDLC), Software Testing Life Cycle
(STLC), Requirements Life Cycle, and Bug Life Cycle methodologies. 
 Proficient in complex, Agile, Scrum, V, and Waterfall methodologies.
 Involved in Requirements Capture, Requirements Categorization, Requirements Prioritization,
Requirements Documentation, Business Rules, stakeholder analysis, Gap analysis and Impact
analysis
 Involved in Sprint meetings and reviews- Entry criteria, Exit criteria, task definition, and
validation criteria.
 Involved in sprint retrospective meetings.
 Good experience in Interface or API testing with Postman, SOAP UI tools, batch testing with Win
SCP, and Notepad++, creating, modifying, and enhancing manual Test cases, and Data migration
applications from legacy systems to As-IS systems.
 Experienced in entire QA life cycle activities, testing concepts, Bug Tracking, Root Cause
Analysis, Defect Analysis, Project Documentation, Designing, Developing, and Generating
Reports, interacting with business analysts, developers, and technical support teams and helping
them in baselining the requirement specifications, maintaining support documents, QA Sign-off
Documents, maintaining Status Reports, creating Requirements Traceability Matrix to ensure
comprehensive test coverage of requirements, preparation of Test plans, Test procedures- and
Manual Test Scenarios. Functional testing of applications, preparing and executing Test Cases
and Test Reports defect reporting, and Test Case Review.
 Coordination and communicated with USA Clients and SA testing efforts, offshore team and
attended Environment, Status, and Defect calls every day to keep the projects online
 Experience in developing SQL queries for backend database testing, and testing applications
migrating from .Net platform to JAVA platform.
TECHNICAL SKILLS:
 Operating Systems: Windows 95, 98, XP/2000/NT, MS-DOS, UNIX.
 Languages: C, C++, C#, JAVA, PEGA
 GUI : VB6.0
 ERP : Mainframes, SAP,
 Database: Oracle10, DB Visualizer, Data Trek
 Web Technologies: VBScript, JavaScript, HTML, VB.Net, ASP. Net, XML, PEGA
 Package: MS-Office
 Testing Tools: QTP12, Quality Center12, LoadRunner11.52
Selenium-Web-Driver, JIRA, SQUIDS, RALLY, ALM, TDD, Microsoft, AZURE, TFS
 CRM : Guidewire, Salesforce, Business Analytics
 Analytics Tools : Rapid Minor
 Diagramming Tools: Axure RP 10
PROFESSIONAL EXPERIENCE:
Franchise Tax Board, State of California Dec’ 2022 – July
2024
Position: Business Analyst/Quality Analyst
Responsibilities:
 Requirements Capture – Engage with business owners to capture Epics and User Stories.
Facilitate workshops.
 Requirements Categorization – Refine and categorize requirements into Functional &amp; Non-
functional requirements.
 Requirements Prioritization – Use MoSCoW methodology to prioritize requirements. Work with
business change owners and IT suppliers to ensure that requirements are clearly understood and
necessary information is available.
 Requirements Documentation – Document requirements catalog, functional specifications, Use
Case Analysis, and UML artifacts.
 Business Rules – Revise and create new business rules for approvals, delegations, and hard-
charging. Conduct stakeholder analysis and identify and resolve
Impact/relationship/dependencies/constraints.
 Ensure strategic objectives, strategy definition, and target models are clear, documented, and
collectively agreed upon.
 Tested PEGA environment and customized the PEGA application as per the client’s
requirements, Integration of PEGA systems with Standalone Forms (SAF), Governance Process
Rules Engine (Rules), Case Management (CM), Return Analysis (RA), Return Verification (RV),
Collections, Notices, CX applications, iCapture, ETS Database and execute the Workflows
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, and User Acceptance Testing, with developers for
requirement development. Tracked Change requests, defects, and test coverage, and published
Defect Status Reports using TFS
 Performed agile activities using the DATABASICS tool. Managed client information for business
communications using Microsoft Dynamics 365
Environment: PEGA, Axure RP 10, DATABASICS, Manual Testing, Report testing API testing,
Microsoft SQL Server, SoapUI, Postman V9.15.2 WINDOWS 10, Google Chrome, Cisco VPN
WINSCP5.15.3, Notepad++, Azure, TFS, Teams, TOSCO.
United Health Group, Child Support State of West Virginia Jun’ 2022 – Dec 2022
Position: Analyst/Tester
Responsibilities:
 Tracking requirements, test coverage using VersionOne tool, defects, published Defect Status
Reports using JIRA
 Created SIT and UAT test cases for the functional testing and regression testing
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, User Acceptance Testing
 Tested Child Support (CSMS), etc modules with sanity flows with migrated data for releasing it
to West Virginia counties.
 Validating Data after data migration, KPI reports data with AS-IS system, the response and result
XMLs, JSON script issues in Postman
 Prepared inbound files and Outbound files. Involved in testing, Batch testing, API interfaces, and
Web Services, analyzing migration issues in Agile methodologies.
 Performed agile activities using the Deltek tool
 Managed client information for business communications using Microsoft Dynamics 365
 Training Maryland county users in using our applications
 Supporting county users when they are facing issues.
 Mentoring the Tester, arranging team meetings, getting updates from each person, forwarding
weekly reports to all the team members and managers, interacting with stakeholders in the
initial stage until handed over automation tester, and forwarding release reports to appropriate
people.
Environment: Sales Force.com, Force.com platform, Service Cloud, Data Loader, ecommerce, ALM,
Rally, TDD, Deltek, Manual Testing, Report testing API testing, Microsoft SQL Server, SoapUI,
Postman V9.15.2 WINDOWS 10, Google Chrome, Cisco VPN WINSCP5.15.3, Notepad++, FACETS
TriZetto.
MDThink, CJAMS State of Maryland Dec’ 2018 – May’ 2022
Position: Analyst/Tester
Responsibilities:
 Tracking requirements, test coverage using VersionOne tool, defects, published Defect Status
Reports using JIRA
 Created SIT and UAT test cases for the functional testing and regression testing
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, User Acceptance Testing
 Tested Child Welfare (CWS), Child Support (CSMS), etc modules in Agile methodology with
sanity flows with migrated data for releasing it to Maryland counties.
 Validating Data after data migration, KPI reports data with AS-IS system, the response and result
XMLs, JSON script issues in Postman
 Prepared inbound files and Outbound files. Involved in testing, Batch testing, API interfaces, and
Web Services, analyzing migration issues
 Performed agile activities using the Digital.ai Agility tool
 Managed client information for business communications using Microsoft Dynamics 365
 Training Maryland county users in using our applications
 Supporting county users when they are facing issues.
 Mentoring the Tester, arranging team meetings, getting updates from each person, forwarding
weekly reports to all the team members and managers, interacting with stakeholders in the
initial stage until handed over automation tester, and forwarding release reports to appropriate
people.
Environment: Sales Force.com, Force.com platform, Service Cloud, Data Loader, ecommerce, JIRA,
VersionOne, Digital.ai Agility, Manual Testing, Report testing API testing, DBeaver /Postgres admin,
SoapUI, Postman V9.15.2 WINDOWS 10, Google Chrome, Open Vpn, Zscaler. WINSCP5.15.3,
Notepad++.
CORVUS, Best Nomos, Senior QA consultant/Team Lead Mar’ 2017 – Dec’2018
Position: Senior QA consultant/Team Lead
Responsibilities:
 Preparation of the Test plan, test approach, Business Requirement document, and deliverables
 Tracking requirements using the DOORS tool, the defects, and published Defect Status Reports
using PTWeb, test coverage using RTM, TTT
 Created test cases for functional testing and regression testing, user acceptance testing using
Excel
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing, Regression
Testing, System Integration Testing, User Acceptance Testing
 Tested CORVUS14, CORVUS Beam Utilities14, CORVUS Configuration Manager14, Corvus User
Manager14, and CORVUS Structure Set Editor modules for Forward planning, Inverse planning,
Treatment planning process in the integration with beam parameters and dose distribution in
modules in Waterfall methodology.
 Mentoring Tester, arranging team meetings, getting updates from each person, forwarding
weekly reports to all the team members and managers, interacting with stakeholders in the
initial stage until handed over automation tester, and forwarding release reports to appropriate
people.
Environment: eCommerce platform, Test Tracking Tool, PTWeb, DOORS, Manual Testing, Mac OS,
Safari, Firefox, ALM and RALLY
SOPHIA, Tele Tech, CO Jan’ 2015 – Feb’2017
Role: Senior QA consultant/Team Lead
Responsibilities:
 Preparation of the Test plan, test approach, and deliverables
 Good experience in creating, modifying, and enhancing manual Test cases
 Created test cases `for the functional testing and regression testing, workflow diagrams for
development references, and data mapping
 Coordinated with Test Analysts to develop Testing scenarios, Functional Testing and Regression
Testing.
 Tested Cards and Payments module for the Card Determination process in the integration with
Customer Management and Payment Type modules in Waterfall methodology.
 Tracked the defects, published Defect Status Reports using ALM, test coverage using RALLY
 Guiding Tester, arranging team meetings, getting updates from each person, forwarding weekly
reports to all the team members and managers, interacting with stakeholders in the initial stage
until handed over automation tester, and forwarding release reports to appropriate people.
Environment: .Net, ALM, and RALLY
Enterprise Pricing Request System (ePRS), FedEx, TN May’2013- Dec’2014
Role: Technical Consultant/Team Lead
Responsibilities:
 Analyzing the requirements from the Business Requirements, Functional Requirements, and
High-level design documents, and test results to make sure requirements have been met and
existing system requirement still holds. 
 Used QC for test management i.e. Defect management, Bug Reporting, and Bug tracking. 
 Coordinating and sending daily status to Project Managers.
 Good experience in creating, modifying, and enhancing manual Test cases in Agile methodology
 Guiding Tester, arranging team meetings, getting updates from each person, forwarding weekly
reports to all the team members and managers, interacting with stakeholders in the initial stage
until handed over to the Client-side Test Manager, and forwarding release reports to
appropriate people.
Environment: E-commerce platform, Java, JSP, ASP servlets, QC12.0
Davis Negative Files, PayPal, San Jose, CA Aug’2012-Apr’2013
Role: Team Lead
Responsibilities:
 Develop test plan, test cases, and test scenarios to map Integration Testing, System Testing, and
User Acceptance Testing to test business requirements and design documents according to the
BRD&#39;s FRDs. Responsible for managing the change request process related to testing
 Involved in execution of SIT and UAT, Cycle 1, Cycle 2 &amp; Cycle 3 Testing, daily meetings to
review the progress of testing as well as discuss the status of defects with offshore people in the
USA. Coordination and communication between US Clients and SA testing Effort.
 Used HP Quality Center (QC) to Report bugs using the defect tracking system and verify fixes for
reserving, executing test cases, defect tracking, and complete test management, to get the
status reports of all the test cases &amp;test scripts that are executed during the testing process, re-
testing of defects corrected and reported.


 Monitor and track priority defects on a day-to-day basis, write positive and negative test cases,
improve traceability, reduce project risk, and increase the quality of the application in Agile
methodology.
Environment: Java, UNIX, JIRA, Java, UNIX, Data Trek.
Retail Staffing Administration, Home Depot, Atlanta, GA Oct’2011-Jul’2012
Role: QA Consultant
Responsibility:
 Develop test cases, and test scenarios to map Integration Testing, System Testing, and User
Acceptance Testing to test business requirements and design documents according to the BRD&#39;s
FRDs. Responsible for managing the change request process related to testing
 Involved in execution of SIT and UAT, Cycle 1, Cycle 2 &amp; Cycle 3 Testing, daily meetings to
review the progress of testing as well as discuss the status of defects with offshore people in the
USA. Coordination and communication between US Clients and SA testing Effort.
 Used HP Quality Center (QC) to Report bugs using the defect tracking system and verify fixes for
reserving, executing test cases, defect tracking, and complete test management, to get the
status reports of all the test cases &amp;test scripts that are executed during the testing process, re-
testing of defects corrected and reported.
 Monitor and track priority defects on a day-to-day basis, write positive and negative test cases,
improve traceability, reduce project risk, and increase the quality of the application in Agile
methodology.
Environment: E-commerce platform, Java, .Net and SOA, QC10, DB Visualizer and QTP
*2007-2016 - SYTEL PVT -Can provide more experience details on request*
EDUCATION:
 MSc. IT/Vinayaka Mission Deemed University-2005
 BSc Computer Science/Osmania University-1998
CERTIFICATIONS:
 GAQM Certified Scrum Master (CSM)®
 SAFe 4 Scrum Master Certificate
 Safe 4 Agilist Certificate
 ISEB ISTQB Certified at the foundation level
 Expert Rating in Software Testing
 Brain bench in Software Testing
 MIT No Code AI and Machine Learning Building Data Science Solutions Certificate
"""

    Format = """
<!DOCTYPE html>
<html lang="en">
<body>

    <h1>Candidate Name</h1>
    <role_title>role title</role_title>

    <h2>PROFESSIONAL SUMMARY</h2>
    <ul>
        <li>professional summary 1</li>
        <li>professional summary 2</li>
        <li>professional summary 3</li>
        <!-- Add more summaries as needed -->
    </ul>


    <h2>TECHNICAL SKILLS</h2>
    <table border="1" cellpadding="5" cellspacing="0">
        <thead>
            <tr>
                <th>Category</th>
                <th>Tools & Technologies</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Category 1</td>
                <td>Tools & Technologies</td>
            </tr>
            <tr>
                <td>Category 2</td>
                <td>Tools & Technologies</td>
            </tr>
            <tr>
                <td>Category 3</td>
                <td>Tools & Technologies</td>
            </tr>
            <!-- Add more rows as needed -->
        </tbody>
    </table>

    <h2>EDUCATION/QUALIFICATION</h2>
    <ul>
        <li>Education 1</li>
        <li>Education 2</li>
        <li>Education 3</li>
        <!-- Add more education qualifications as needed -->
    </ul>

    <h2>CERTIFICATION/TRAINING</h2>
    <ul>
        <li>Certification 1</li>
        <li>Certification 2</li>
        <li>Certification 3</li>
        <!-- Add more certifications as needed -->
    </ul>


    <!-- Strict instruction: Use the format MM/YY - MM/YY (Total Months) and calculate the total months & set 'Present' if applicable. If the candidate is currently working, ensure all responsibilities are written in the present tense. for all of the past experiences of candidates, ensure all responsibilities are written in the past tense. Double-check for consistent tense usage across all responsibilities. -->

    <h2>WORK HISTORY</h2>

    <strong class='Add "Previous experience" or "Present experience" in class based on provide reference date' >Date: MM/YY - MM/YY (Total Months)</strong> 
    <strong>Company: Company Name</strong>

        <p><strong>Client:</strong> Client Name</p>
        <p><strong>Title:</strong> Job Title</p>
        <p><strong>Tools and Technologies:</strong> Tools, Technologies</p>
        <p><strong>Description:</strong> Job description goes here</p>

        <p><strong>Roles and Responsibilities</strong></p> 
        <ol>
            <li>Responsibility 1</li> 
            <li>Responsibility 2</li> 
            <li>Responsibility 3</li> 
            <!-- Add more responsibilities as needed -->
        </ol>


    <!-- Repeat the WORK HISTORY section for each job -->

</body>
</html>
"""

    keywords = []

    messages = [
        {
            "role": "system",
            "content": """
                                You are an expert in resume formatting, with a specialty in ensuring adherence to predefined templates with high precision. Your task involves accurately parsing and reformatting resumes according to a specified structure. All information must be captured and presented in an organized, HTML-compatible format, while maintaining grammatical accuracy. Your focus is on ensuring no unnecessary details are included and all required components are present.
                            """,
        },
        {
            "role": "user",
            "content": f"""
                        You are provided with the following resume data to process:

                        Instructions:
                        Professional Summary: Create a concise, well-crafted summary of the candidate's skills and accomplishments.
                        Technical Skills: Present the candidate's skills in a tabular format for clarity.
                        Work History: For each position, provide the dates, company, title, environment, job description, and responsibilities. For ongoing roles, ensure details (Roles and Responsibilities) are written in the present tense; for past roles, details (Roles and Responsibilities) must be written in the past tense. And also make sure that all information is complete and no details are omitted or missed out.
                        Provide detailed descriptions for each project, ensuring all responsibilities and roles are fully included without summarization or omission.
                        Ensure the tense is appropriate for each role:
                        For ongoing roles: use present tense (e.g., “Developing and maintaining web applications”).
                        For completed roles: use past tense (e.g., “Developed and maintained web applications”).
                        Ensure all responsibilities are grammatically correct and that the correct tense is applied consistently throughout.
                        Projects: For each project, provide the client, designation, environment, description, and responsibilities. Check if any of the client, designation, environment, description, are mentioned in the project details, then keep it blank
                        Education: Include all relevant educational qualifications in a structured format.
                        Certifications/Training: List any certifications or training completed.
                        Consistency Check: After processing, thoroughly review the resume:
                        Verify that ongoing roles use only present tense.
                        Verify that past roles use only past tense.
                        Ensure there are no grammatical inconsistencies or tense errors.
                        Keyword Highlighting: Bold the words from the provided list by enclosing each instance of those words in <b> tags throughout the resume. Only these specified words should be bolded. But not make it case sensitive.
                        Formatting Guidelines:
                        Use <h2> tags for section headings.
                        Use <p> tags for descriptions.
                        Use <ul> or <ol> and <li> tags for bullet-pointed lists.
                        Ensure the format is clear and does not include unnecessary information.
                        Consistently apply bold formatting to the words from the provided list using <b> tags. Only those words should be bolded.
                        Validation and Error Handling:
                        Perform a manual and automated review of the resume to ensure grammatical precision, tense consistency, and correct formatting.
                        Immediately address any detected errors and revalidate to ensure 100% accuracy.

                        Giving a reference of the one of the resume

                        Original Resume: [{unformatted_resume}]
                        Formatted Resume: [{formatted_resume}]

                        Resume: [{content}]
                        Format: [{Format}]
                        Word List: [{keywords}]
                        """,
        },
    ]

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.15,
        n=1,
    )

    reply = completion.choices[0].message.content

    if "```html" in reply and "```" in reply:
        reply = reply.replace("```html", "").replace("```", "").strip()
    if keywords:
        return wrap_keywords_in_b_tags(reply, keywords)
    return reply


# /////////////////////////// DOCX to HTML ///////////////////////////


def add_header_with_logo_and_contact(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Create table with 1 row and 2 columns
        table = header.add_table(1, 2, width=Cm(20.32))
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.columns[0].width = Cm(7.62)  # Adjust column width for the logo
        table.columns[1].width = Cm(12.7)  # Adjust column width for the contact info

        # Left cell for logo
        left_cell = table.cell(0, 0)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if os.path.exists(LOGO_PATH):
            paragraph = left_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(
                LOGO_PATH, width=Cm(4.67), height=Cm(2.3)
            )  # Adjust logo size
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            print(f"Logo file not found at {LOGO_PATH}")

        # Right cell for contact info
        right_cell = table.cell(0, 1)
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        contact_paragraph = right_cell.paragraphs[0]
        contact_run = contact_paragraph.add_run(CONTACT_INFO)
        contact_run.font.size = Pt(10)
        contact_run.font.name = FONT_NAME
        contact_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )  # Align text to the right

    # Add a line break after the logo
    doc.add_paragraph()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Remove extra space before/after paragraphs
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)


def add_paragraph(doc, text, style=None, bold=False, alignment=None, color=None):
    p = doc.add_paragraph(text, style=style)
    if alignment:
        p.alignment = alignment
    run = p.runs[0]
    run.bold = bold
    run.font.color.rgb = color if color else RGBColor(0, 0, 0)  # Set text color
    run.font.name = FONT_NAME  # Set font to Times New Roman
    run.font.size = Pt(10)  # Set font size to 10


def add_list_item(doc, element, indent):
    p = doc.add_paragraph(style="List Bullet")
    for child in element.children:
        handle_element(doc, child, p)
    p.paragraph_format.left_indent = Cm(indent)  # Adjust the indent as needed


def handle_element(doc, element, parent_paragraph=None):
    if isinstance(element, str):
        if parent_paragraph:
            run = parent_paragraph.add_run(element)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        return

    # Handle <b> and <strong> tags for bold text
    if element.name in ["b", "strong"]:
        if parent_paragraph:
            run = parent_paragraph.add_run(element.get_text())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)  # Set font size to 10
            run.font.color.rgb = RGBColor(0, 0, 0)
            if element.name == "strong":
                parent_paragraph.paragraph_format.left_indent = Cm(0.64)
                parent_paragraph.paragraph_format.space_before = Pt(12)
                parent_paragraph.paragraph_format.space_after = Pt(12)
        else:
            add_paragraph(doc, element.get_text(), bold=True)
        return

    if element.name == "h1":
        add_paragraph(
            doc, element.get_text(), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
    elif element.name == "role_title":
        add_paragraph(
            doc,
            element.get_text(),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            color=RGBColor(0, 0, 255),
        )
    elif element.name == "h2":
        p = doc.add_paragraph(element.get_text())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)  # Add space before h2
        p_format.space_after = Pt(12)  # Add space after h2
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
        run.font.name = FONT_NAME  # Set font to Times New Roman
        run.font.size = Pt(10)  # Set font size to 10
    elif element.name == "h3":
        add_paragraph(doc, element.get_text(), style="Heading 2")
    elif element.name == "h4":
        add_paragraph(doc, element.get_text(), style="Heading 3")
    elif element.name == "p":
        p = doc.add_paragraph()
        for child in element.children:
            handle_element(doc, child, p)
    elif element.name == "table":
        # Create table with a specific number of columns (based on <th> tags)
        table = doc.add_table(
            rows=1, cols=len(element.find_all("th")), style="Table Grid"
        )

        # Disable autofit to allow manual control over the column widths
        table.autofit = False

        # Set the column widths (e.g., 1 inch for each column, adjust as needed)
        for column in table.columns:
            for cell in column.cells:
                cell.width = Cm(2.54)  # Adjust column width as needed

        # Set header row with custom font and size
        hdr_cells = table.rows[0].cells
        for idx, th in enumerate(element.find_all("th")):
            hdr_cells[idx].text = th.get_text()
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Set font size to 10pt
                    run.font.name = FONT_NAME  # Set font to Times New Roman

        # Add table rows and handle table data
        for tr in element.find_all("tr")[1:]:
            row_cells = table.add_row().cells
            for idx, td in enumerate(tr.find_all("td")):
                cell_paragraph = row_cells[idx].paragraphs[0]
                handle_element(doc, td, cell_paragraph)

                # Align paragraph and vertically center content in each cell
                row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Apply indentation to the table (e.g., 1 inch indent)
        tbl = table._tbl  # Access the underlying table element
        tblPr = tbl.tblPr  # Get the table properties
        tblInd = OxmlElement("w:tblInd")  # Create table indentation element
        tblInd.set(qn("w:w"), "500")  # Set indentation value in twips (1440 = 1 inch)
        tblInd.set(qn("w:type"), "dxa")  # Set measurement type to dxa (twips)

        # Append the indentation element to the table properties
        tblPr.append(tblInd)
    elif element.name in ["ul", "ol"]:
        indent = 1.27
        if element.name == "ol":
            indent = 2.12
        for li in element.find_all("li"):
            add_list_item(doc, li, indent)
        # Add space after the list
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    elif element.name == "br":
        doc.add_paragraph()
    else:
        for child in element.children:
            handle_element(doc, child, parent_paragraph)


def convert_html_to_docx(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()

    add_header_with_logo_and_contact(doc)

    for element in soup.body:
        handle_element(doc, element)

    return doc
    # doc.save(docx_filename)


# //////////////////////////////////////////////////////////////////////
def read_pdf(file_path):
    """Reads a .pdf file and returns its content as a string.

    Args:
        file_path (str): The path to the PDF file.

    Returns:
        str: The content of the PDF file as a string.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")

    try:
        # Open the PDF file
        doc = fitz.open(file_path)
        content = []

        # Iterate through each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            content.append(page.get_text())

        return "\n".join(content)

    except Exception as e:
        raise RuntimeError(f"An error occurred while reading the PDF file: {e}")


def convert_doc_to_docx(doc_path, docx_path):
    """Converts a .doc file to .docx.

    Args:
        doc_path (str): The path to the .doc file.
        docx_path (str): The path to save the .docx file.
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"File not found: {doc_path}")

    word = win32com.client.Dispatch("Word.Application")

    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, FileFormat=16)  # FileFormat=16 for .docx
        doc.Close()
    except Exception as e:
        raise RuntimeError(f"An error occurred while converting .doc to .docx: {e}")
    finally:
        word.Quit()


def convert_and_read(file_path):
    """Converts .doc or .docx files to .pdf and reads the .pdf content, or reads .pdf content directly.

    Args:
        file_path (str): The path to the input file (.doc, .docx, or .pdf).

    Returns:
        str: The content of the PDF file as a string.
    """
    # Handle .doc files by converting them to .docx
    if file_path.endswith(".doc"):
        docx_path = file_path.replace(".doc", ".docx")
        convert_doc_to_docx(file_path, docx_path)
        file_path = docx_path

    # Convert .docx files to .pdf
    if file_path.endswith(".docx"):
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert(file_path, pdf_path)
        except Exception as e:
            raise RuntimeError(f"An error occurred during the conversion: {e}")
        file_path = pdf_path

    # Handle .pdf files
    if file_path.endswith(".pdf"):
        return read_pdf(file_path)

    raise ValueError(
        "Unsupported file type. Please provide a .doc, .docx, or .pdf file."
    )


def handle_temp_file(uploaded_file, suffix):

    # Initialize COM environment
    pythoncom.CoInitialize()

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name

        # Perform the conversion
        resume = convert_and_read(tmp_file_path)

        # Cleanup
        os.unlink(tmp_file_path)
        return resume
    finally:
        # Ensure COM environment is uninitialized
        pythoncom.CoUninitialize()


# Streamlit UI

st.title("Enhanced Resume Reformatter")
st.write(
    "Upload a resume in DOC,DOCX or PDF format to convert it to the predefined format."
)

api_key = st.text_input("Enter your OpenAI API key:", type="password")
keywords_input = st.text_input("Enter keywords (comma-separated)")
uploaded_file = st.file_uploader("Choose a file", type=["docx", "pdf"])
keywords = [keyword.strip() for keyword in keywords_input.split(",") if keyword.strip()]

if uploaded_file is not None:
    try:
        client = initialize_openai_client(api_key)
        file_type = uploaded_file.type

        if (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            resume_content = handle_temp_file(uploaded_file, ".docx")

        elif file_type == "application/msword":
            resume_content = handle_temp_file(uploaded_file, ".doc")

        elif file_type == "application/pdf":
            file_content = uploaded_file.getvalue()
            resume_content = extract_content_from_pdf(io.BytesIO(file_content))
            st.write("PDF file uploaded successfully.")

        else:
            st.error("Unsupported file format. Please upload a PDF, DOC, or DOCX file.")
            st.stop()

        cleaned_resume_content = clean_text(resume_content)

        with st.spinner("Reformatting resume..."):
            formatted_resume = reformat_resume(cleaned_resume_content, keywords)

        final_formatted_doc = convert_html_to_docx(formatted_resume)

        # Save as DOCX
        docx_buffer = io.BytesIO()
        final_formatted_doc.save(docx_buffer)
        docx_buffer.seek(0)

        st.download_button(
            label="Download Final Formatted Resume (DOCX)",
            data=docx_buffer,
            file_name="Final_Formatted_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        st.error("Please try uploading the file again or use a different file.")
