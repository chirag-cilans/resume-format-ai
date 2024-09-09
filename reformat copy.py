import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
from openai import OpenAI
import tempfile
import os
import subprocess

# Initialize the OpenAI client
def initialize_openai_client(api_key):
    return OpenAI(api_key=api_key)
def extract_content_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path], check=True)
        return pdf_path
    except subprocess.CalledProcessError as e:
        raise Exception(f"Error converting DOCX to PDF: {e}")

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = ''.join(char for char in text if char.isprintable() or char in ['\n', '\t'])
    return text.strip()

def reformat_resume(content):
    prompt = f"""
You are an expert resume formatter with a perfect track record. Your task is to reformat the provided resume content into a standardized structure. Follow these instructions meticulously:

1. Analyze the entire input content thoroughly before starting the formatting process.
2. Ensure that ALL information from the input is included in the output, without exception. Do not add any information that is not present in the original content.
3. Format the resume EXACTLY according to the structure provided below, populating all sections completely.
4. Do not use any placeholders, continuation instructions, or omit any information.
5. Do not add any additional headings or subheadings that are not specified in the output format.
6. If a section seems empty based on the input, include the section header and state "No information provided" underneath.
7. For the WORK HISTORY section, include ALL positions from the input content, no matter how many there are.
8. Calculate and include the duration for each work history entry.
9. Maintain professional language and correct any obvious typos or grammatical errors in the content.
10. Bold only the most important words or phrases in each section, especially in the PROFESSIONAL SUMMARY and WORK HISTORY sections. Be selective with bolding.
11. For TECHNICAL SKILLS, provide a simple comma-separated list of up to 15 skills. If more than 15 skills are present, choose the most important ones. Do not categorize or group the skills.
12. Ensure that the PROFESSIONAL SUMMARY is populated with relevant content from the input, but keep it concise (2-3 sentences maximum).
13. Make sure the Position Title is included right after the Full Name.
14. In the WORK HISTORY section, limit the number of bullet points under roles and responsibilities to a maximum of 6 for each position. Choose the most important responsibilities if there are more than 6.

Input Resume Content:
{content}

Output Format:

[Full Name]
[Position Title]

PROFESSIONAL SUMMARY
[Generate a concise professional summary (2-3 sentences) based on the provided content, highlighting key skills and experience. Bold only the most important words or phrases.]

TECHNICAL SKILLS
[List up to 15 most important technical skills mentioned in the input, separated by commas, without any categorization or grouping]

EDUCATION QUALIFICATIONS
[List all educational qualifications in the format: Degree, Institution, Year]

CERTIFICATION/TRAINING
[List all certifications and training mentioned in the input]

WORK HISTORY
[For each position, use the following format. Include ALL positions from the input content:]
Date: [Start Date] - [End Date] (Duration: [Calculate and insert duration])
Company: [Company Name], [Location (if provided)]
Title: [Position Title]
Roles and Responsibilities:
- [Responsibility 1 with only the most important words or phrases in bold]
- [Responsibility 2 with only the most important words or phrases in bold]
- [Responsibility 3 with only the most important words or phrases in bold]
[Include up to 6 most important responsibilities for each position]

[Repeat the above format for each position in the work history, ensuring no positions are omitted]

Ensure that every section is populated with all relevant information from the input. Do not omit any details or use placeholders. Do not add any sections or subheadings not specified above. Format the entire resume completely and professionally.
    """
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a meticulous resume formatting assistant. Your task is to format the entire resume completely and accurately, ensuring all information is included and properly structured according to the exact output format specified."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000,
        n=1,
        temperature=0.2,
    )
    return response.choices[0].message.content.strip()

def add_section_header(doc, text):
    doc.add_paragraph()
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.space_after = Pt(6)

    doc.add_paragraph()

def add_content(doc, text, indent=False):
    paragraph = doc.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Pt(18)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_bullet_point(doc, text):
    paragraph = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

def create_skills_table(doc, skills):
    total_skills = len(skills)
    cols = 3
    rows = -(-total_skills // cols)  # Ceiling division

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = False

    for i, skill in enumerate(skills):
        row = i // cols
        col = i % cols
        cell = table.cell(row, col)
        cell.text = skill.strip()
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    for column in table.columns:
        column.width = Inches(6.0 / cols)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            set_cell_border(cell)

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_header_with_logo_and_contact(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        table = header.add_table(1, 2, width=Inches(8))
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        left_cell = table.cell(0, 0)
        left_cell.width = Inches(4)
        logo_path = os.path.join(os.path.dirname(__file__), "kyralogo.png")
        if os.path.exists(logo_path):
            paragraph = left_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(2))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            print(f"Logo file not found at {logo_path}")
        
        right_cell = table.cell(0, 1)
        right_cell.width = Inches(4)
        contact_info = """3673 Coolidge Ct., Tallahassee, FL 32311
Phone: (850) 459-5854
Email: vpatel@KyraSolutions.com"""
        contact_paragraph = right_cell.paragraphs[0]
        contact_run = contact_paragraph.add_run(contact_info)
        contact_run.font.size = Pt(8)
        contact_run.font.name = 'Times New Roman'
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add a line break after the logo
    doc.add_paragraph()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

def save_to_word(formatted_content):
    doc = Document()
    
    add_header_with_logo_and_contact(doc)
    
    sections = formatted_content.split('\n\n')
    
    for index, section in enumerate(sections):
        lines = section.split('\n')
        if not lines:
            continue
        
        section_title = lines[0].strip()
        
        if index == 0:  # This is the name and position
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(lines[0])  # Name
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if len(lines) > 1:  # Position title
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(lines[1])
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
        else:
            if section_title == "PROFESSIONAL SUMMARY":
                add_section_header(doc, section_title)
                for line in lines[1:]:
                    add_content(doc, line)
            elif section_title == "TECHNICAL SKILLS":
                add_section_header(doc, section_title)
                skills = [skill.strip() for skill in ' '.join(lines[1:]).split(',')]
                create_skills_table(doc, skills)
            else:
                add_section_header(doc, section_title)
                
                if section_title == "WORK HISTORY":
                    job_entries = section.split('Date:')[1:]
                    for job in job_entries:
                        job_lines = ['Date:' + job.strip()][0].split('\n')
                        for line in job_lines:
                            if line.startswith('Date:'):
                                add_content(doc, line, indent=True)
                            elif line.startswith('Company:'):
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(line)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
                                paragraph.paragraph_format.left_indent = Pt(18)
                            elif line.startswith('-'):
                                add_bullet_point(doc, line[2:])
                            elif any(line.startswith(prefix) for prefix in ['Title:', 'Roles and Responsibilities:']):
                                add_content(doc, line, indent=True)
                            elif line.strip():
                                add_content(doc, line)
                        doc.add_paragraph()
                else:
                    for line in lines[1:]:
                        if line.startswith('-'):
                            add_bullet_point(doc, line[2:])
                        else:
                            add_content(doc, line)
    
    return doc

# Streamlit UI

st.title('Enhanced Resume Reformatter')
st.write('Upload a resume in DOCX or PDF format to convert it to the predefined format.')

api_key = st.text_input("Enter your OpenAI API key:", type="password")
uploaded_file = st.file_uploader("Choose a file", type=["docx", "pdf"])

if uploaded_file is not None:
    try:
        client = initialize_openai_client(api_key)
        file_type = uploaded_file.type
        
        if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Convert DOCX to PDF
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                tmp_docx.write(uploaded_file.getvalue())
                tmp_docx_path = tmp_docx.name

            pdf_path = convert_docx_to_pdf(tmp_docx_path)
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()

            os.unlink(tmp_docx_path)
            os.unlink(pdf_path)
        elif file_type == "application/pdf":
            pdf_content = uploaded_file.getvalue()
        else:
            st.error("Unsupported file format. Please upload a PDF or DOCX file.")
            st.stop()

        # Extract content from PDF
        resume_content = extract_content_from_pdf(io.BytesIO(pdf_content))
        cleaned_resume_content = clean_text(resume_content)

        with st.spinner('Reformatting resume...'):
            formatted_resume = reformat_resume(cleaned_resume_content)

        sections = formatted_resume.split('\n\n')
        edited_sections = []

        st.subheader("Edit Formatted Resume Content:")
        for i, section in enumerate(sections):
            if section.strip():
                section_title = section.split('\n')[0]
                section_content = '\n'.join(section.split('\n')[1:])
                edited_content = st.text_area(
                    f"{section_title}:", 
                    section_content, 
                    height=200, 
                    key=f"section_{i}"
                )
                edited_sections.append(f"{section_title}\n{edited_content}")

        final_formatted_resume = '\n\n'.join(edited_sections)

        if st.button("Generate Final Resume"):
            final_formatted_doc = save_to_word(final_formatted_resume)

            # Save as DOCX
            docx_buffer = io.BytesIO()
            final_formatted_doc.save(docx_buffer)
            docx_buffer.seek(0)

            st.download_button(
                label="Download Final Formatted Resume (DOCX)",
                data=docx_buffer,
                file_name="Final_Formatted_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        st.error("Please try uploading the file again or use a different file.")